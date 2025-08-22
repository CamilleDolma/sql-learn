from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class WordGenerator:
    """
    一个用于生成符合特定格式要求的 Word 文档的类。

    方法:
        add_title(text): 添加文档标题。
        add_abstract(text): 添加摘要。
        add_keywords(keywords): 添加关键词。
        add_body_paragraph(text): 添加正文段落。
        add_references(references): 添加参考文献。
        save(): 保存文档。
    """
    def __init__(self, output_filename="成品文档.docx"):
        """
        初始化 WordGenerator。

        Args:
            output_filename (str): 输出的 Word 文档文件名。
        """
        self.doc = Document()
        self.output_filename = output_filename
        self._set_page_setup()

    def _set_page_setup(self):
        """设置页面为 A4 大小，边距使用 Word 默认值。"""
        section = self.doc.sections[0]
        # A4 paper size in points
        section.page_width = Pt(595.3)
        section.page_height = Pt(841.9)
        # 使用 Word 默认页边距

    def _set_paragraph_style(self, paragraph, font_name, font_size_pt, is_bold=False, 
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None):
        """
        一个通用的段落样式设置函数。

        Args:
            paragraph: docx 中的 paragraph 对象。
            font_name (str): 字体名称 (例如 '黑体')。
            font_size_pt (int): 字体大小 (磅)。
            is_bold (bool): 是否加粗。
            alignment: 对齐方式。
            first_line_indent: 首行缩进值。
        """
        paragraph.alignment = alignment
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(24)  # 固定行距24磅
        if first_line_indent:
            paragraph_format.first_line_indent = first_line_indent

        # 确保段落中至少有一个 run
        if not paragraph.runs:
            paragraph.add_run()
        
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font.bold = is_bold
            font.size = Pt(font_size_pt)
            # 设置中文字体
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def add_title(self, text):
        """添加题目 (2号黑体, 居中)。"""
        p = self.doc.add_paragraph(text)
        self._set_paragraph_style(p, '黑体', 22, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def add_abstract(self, text):
        """添加摘要部分 (小4号仿宋体)。"""
        heading = self.doc.add_paragraph('摘要')
        self._set_paragraph_style(heading, '黑体', 12, is_bold=True) # 摘要标题使用小四黑体

        content = self.doc.add_paragraph(text)
        # 正文首行缩进2个字符 (小四字体约24磅)
        self._set_paragraph_style(content, '仿宋', 12, first_line_indent=Pt(24))

    def add_keywords(self, keywords):
        """添加关键词部分 (小4号仿宋体)。"""
        heading = self.doc.add_paragraph('关键词')
        self._set_paragraph_style(heading, '黑体', 12, is_bold=True) # 关键词标题使用小四黑体

        content_text = '；'.join(keywords)
        content = self.doc.add_paragraph(content_text)
        self._set_paragraph_style(content, '仿宋', 12)

    def add_body_paragraph(self, text):
        """添加正文段落 (小4号仿宋体, 首行缩进)。"""
        p = self.doc.add_paragraph(text)
        self._set_paragraph_style(p, '仿宋', 12, first_line_indent=Pt(24))

    def add_references(self, references):
        """添加参考文献 (小4号仿宋体)。"""
        heading = self.doc.add_paragraph('参考文献')
        self._set_paragraph_style(heading, '黑体', 12, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        for ref_text in references:
            p = self.doc.add_paragraph(ref_text)
            self._set_paragraph_style(p, '仿宋', 12)

    def save(self):
        """保存 Word 文档。"""
        self.doc.save(self.output_filename)
        print(f"文档 '{self.output_filename}' 已成功生成。")


if __name__ == '__main__':
    # --- API 使用示例 ---
    doc_gen = WordGenerator("我的毕业论文.docx")

    # 1. 添加题目
    doc_gen.add_title("关于人工智能在文学创作中的应用研究")

    # 2. 添加摘要
    doc_gen.add_abstract("本文探讨了人工智能技术在文学创作领域的应用现状、挑战与未来发展趋势。通过分析当前的AI写作工具...")

    # 3. 添加关键词
    doc_gen.add_keywords(["人工智能", "文学创作", "AI写作", "计算语言学"])

    # 4. 添加正文段落 (可以多次调用)
    doc_gen.add_body_paragraph("第一章 引言。随着深度学习技术的飞速发展，人工智能已经渗透到社会生活的方方面面...")
    doc_gen.add_body_paragraph("第二章 相关技术综述。本章将详细介绍用于文本生成的几种主流AI模型，包括循环神经网络（RNN）、长短期记忆网络（LSTM）以及最新的Transformer架构...")

    # 5. 添加参考文献
    doc_gen.add_references([
        "[1] 张三. 人工智能导论[M]. 北京: 高等教育出版社, 2020.",
        "[2] 李四. 基于生成对抗网络的诗歌生成研究[D]. 南京: 南京大学, 2021.",
        "[3] Wang L, Li W. A survey of recent advances in neural text generation[J]. Journal of Computer Science and Technology, 2022, 37(1): 1-20."
    ])

    # 6. 保存文档
    doc_gen.save()